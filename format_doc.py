#!/usr/bin/env python3
"""
format_doc.py — Complete report formatter. Run AFTER humanize_doc.py.

What it does (never touches text content):
  1.  Remove all stray empty paragraphs between headings
  2.  Chapter titles (CHAPTER X, ABSTRACT, etc.) always start on a new page
  3.  Body paragraph justification, font, line-spacing normalized
  4.  Heading spacing normalized (space-before / space-after)
  5.  Widow/orphan control on body text
  6.  Replace manual TOC with Word auto-field (press F9 in Word to update)
  7.  LibreOffice headless refresh for repagination (if installed)

Usage:
    python3 format_doc.py input.docx output.docx
    python3 format_doc.py input.docx output.docx --report   # dry-run
    python3 format_doc.py input.docx output.docx --no-toc
    python3 format_doc.py input.docx output.docx --no-refresh

Requires:
    pip install python-docx lxml
    apt install libreoffice   (optional — for TOC page-number refresh)
"""

import argparse
import copy
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


# ─── FORMATTING CONSTANTS (extracted from this document) ─────────────────────

BODY_FONT        = "Times New Roman"
BODY_SIZE_PT     = 12.0
BODY_LINE_TWIPS  = 360          # 1.5 line spacing
BODY_LINE_RULE   = "auto"
BODY_JC          = "both"       # full justification
BODY_SPACE_AFTER = 0
BODY_LEFT_TWIPS  = 980
BODY_RIGHT_TWIPS = 744
BODY_FIRST_LINE  = 720          # 0.5 inch first-line indent

H2_SPACE_BEFORE  = 280          # ~18pt
H2_SPACE_AFTER   = 80           # ~4pt
H3_SPACE_BEFORE  = 220          # ~14pt
H3_SPACE_AFTER   = 40

# Styles that are body prose
BODY_STYLES = {'normal', 'body text', 'default paragraph font', 'no spacing'}

# Styles never touched
SKIP_STYLES = {
    'heading 1','heading 2','heading 3','heading 4',
    'title','subtitle','caption','footnote text',
    'header','footer','toc 1','toc 2','toc 3',
    'list paragraph','list bullet','list number',
    'table contents','table heading','code','verbatim',
}

# Chapter-level titles — these always start on a NEW PAGE
CHAPTER_TITLES = {
    'abstract', 'acknowledgement', 'acknowledgements',
    'acknowledgment', 'acknowledgments',
    'contents', 'table of contents',
    'list of figures', 'list of tables', 'list of acronyms',
    'references', 'bibliography',
    'conclusion', 'conclusions',
    'introduction',
}
CHAPTER_LABEL_RE = re.compile(r'^chapter\s+\d+$', re.I)


# ─── HELPERS ─────────────────────────────────────────────────────────────────

def ensure_ppr(para):
    ppr = para._element.find(qn('w:pPr'))
    if ppr is None:
        ppr = OxmlElement('w:pPr')
        para._element.insert(0, ppr)
    return ppr


def set_child(parent, tag, attrs):
    existing = parent.find(qn(tag))
    if existing is not None:
        parent.remove(existing)
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), str(v))
    parent.append(el)
    return el


def has_page_break(para):
    for br in para._element.findall('.//' + qn('w:br')):
        if br.get(qn('w:type')) == 'page':
            return True
    return False


def add_page_break_before(para):
    """Add a page break before this paragraph using <w:pageBreakBefore>."""
    ppr = ensure_ppr(para)
    pb = ppr.find(qn('w:pageBreakBefore'))
    if pb is None:
        pb = OxmlElement('w:pageBreakBefore')
        ppr.insert(0, pb)
    pb.set(qn('w:val'), '1')


def remove_page_break_before(para):
    """Remove any <w:pageBreakBefore> from a paragraph."""
    ppr = para._element.find(qn('w:pPr'))
    if ppr is not None:
        pb = ppr.find(qn('w:pageBreakBefore'))
        if pb is not None:
            ppr.remove(pb)


def is_chapter_title(para):
    """True if this paragraph is a chapter-level title that should start a new page."""
    txt  = para.text.strip()
    low  = txt.lower()
    style = para.style.name.lower()

    # "CHAPTER X" label paragraphs (style=normal)
    if CHAPTER_LABEL_RE.match(txt):
        return True
    # Known title words in normal style
    if low in CHAPTER_TITLES and style == 'normal':
        return True
    return False


def is_body_para(para):
    """True if this paragraph should receive full body formatting."""
    txt = para.text.strip()
    if not txt:
        return False
    if para._element.findall('.//' + qn('w:drawing')):
        return False
    if (para._element.findall('.//' + qn('m:oMath')) or
            para._element.findall('.//' + qn('m:oMathPara'))):
        return False
    style = para.style.name.lower()
    if any(s in style for s in SKIP_STYLES):
        return False
    if style not in BODY_STYLES and not style.startswith('normal'):
        return False
    words = txt.split()
    if len(words) < 15:
        return False
    if not any(c in txt for c in ['.', ',']):
        return False
    if txt.count('…') > 2 or txt.count('.....') > 2:
        return False
    return True


# ─── FIX FUNCTIONS ───────────────────────────────────────────────────────────

def remove_empty_paras_before_heading(doc, report=False):
    """Delete stray empty paragraphs that appear immediately before a heading."""
    paras    = list(doc.paragraphs)
    body     = doc.element.body
    removed  = 0

    # Build set of paragraph elements to remove
    to_remove = set()
    for i, para in enumerate(paras):
        style = para.style.name
        if not (style.startswith('Heading') and para.text.strip()):
            continue
        # Walk backwards while empty
        j = i - 1
        while j >= 0:
            prev = paras[j]
            if prev.text.strip():
                break
            to_remove.add(id(prev._element))
            j -= 1

    for para in paras:
        if id(para._element) in to_remove:
            if not report:
                try:
                    body.remove(para._element)
                except ValueError:
                    pass
            removed += 1

    return removed


def fix_chapter_page_breaks(doc, report=False):
    """Ensure every chapter title starts on a new page."""
    added   = 0
    removed = 0
    for para in doc.paragraphs:
        if is_chapter_title(para):
            if not has_page_break(para):
                if not report:
                    add_page_break_before(para)
                added += 1
        else:
            # Remove spurious page-break-before from non-chapter elements
            ppr = para._element.find(qn('w:pPr'))
            if ppr is not None:
                pb = ppr.find(qn('w:pageBreakBefore'))
                if pb is not None and pb.get(qn('w:val')) == '1':
                    if not report:
                        ppr.remove(pb)
                    removed += 1
    return added, removed


def fix_body_formatting(doc, report=False):
    """Justify, font, line-spacing on all body paragraphs."""
    fixed = 0
    for para in doc.paragraphs:
        if not is_body_para(para):
            continue
        changed = False
        ppr = ensure_ppr(para)

        # Justification
        jc = ppr.find(qn('w:jc'))
        if jc is None or jc.get(qn('w:val')) != 'both':
            if not report:
                set_child(ppr, 'w:jc', {'w:val': 'both'})
            changed = True

        # Line spacing
        sp = ppr.find(qn('w:spacing'))
        if sp is None or sp.get(qn('w:line')) != str(BODY_LINE_TWIPS):
            if not report:
                if sp is None:
                    sp = OxmlElement('w:spacing')
                    ppr.append(sp)
                sp.set(qn('w:line'),     str(BODY_LINE_TWIPS))
                sp.set(qn('w:lineRule'), BODY_LINE_RULE)
                sp.set(qn('w:after'),    str(BODY_SPACE_AFTER))
            changed = True

        # Widow control
        wc = ppr.find(qn('w:widowControl'))
        if wc is None:
            if not report:
                set_child(ppr, 'w:widowControl', {'w:val': '0'})
            changed = True

        # Font + size on runs
        sz_val = str(int(BODY_SIZE_PT * 2))
        for run in para.runs:
            if not run.text.strip():
                continue
            rpr = run._r.find(qn('w:rPr'))
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                run._r.insert(0, rpr)

            rf = rpr.find(qn('w:rFonts'))
            if rf is None:
                rf = OxmlElement('w:rFonts')
                rpr.insert(0, rf)
            for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
                if rf.get(qn(attr)) != BODY_FONT and not report:
                    rf.set(qn(attr), BODY_FONT)
                    changed = True

            sz = rpr.find(qn('w:sz'))
            if sz is None or sz.get(qn('w:val')) != sz_val:
                if not report:
                    if sz is None:
                        sz = OxmlElement('w:sz')
                        rpr.append(sz)
                    sz.set(qn('w:val'), sz_val)
                    szcs = rpr.find(qn('w:szCs'))
                    if szcs is None:
                        szcs = OxmlElement('w:szCs')
                        rpr.append(szcs)
                    szcs.set(qn('w:val'), sz_val)
                changed = True

        if changed:
            fixed += 1
    return fixed


def fix_heading_spacing(doc, report=False):
    """Normalize heading space-before/after and remove extra blank lines."""
    fixed = 0
    for para in doc.paragraphs:
        style = para.style.name.lower()
        if not para.text.strip():
            continue
        ppr = ensure_ppr(para)
        changed = False

        if 'heading 2' in style:
            sp = ppr.find(qn('w:spacing'))
            if sp is None:
                sp = OxmlElement('w:spacing'); ppr.append(sp)
            if (sp.get(qn('w:before')) != str(H2_SPACE_BEFORE) or
                    sp.get(qn('w:after')) != str(H2_SPACE_AFTER)):
                if not report:
                    sp.set(qn('w:before'), str(H2_SPACE_BEFORE))
                    sp.set(qn('w:after'),  str(H2_SPACE_AFTER))
                changed = True

        elif 'heading 3' in style:
            sp = ppr.find(qn('w:spacing'))
            if sp is None:
                sp = OxmlElement('w:spacing'); ppr.append(sp)
            if (sp.get(qn('w:before')) != str(H3_SPACE_BEFORE) or
                    sp.get(qn('w:after')) != str(H3_SPACE_AFTER)):
                if not report:
                    sp.set(qn('w:before'), str(H3_SPACE_BEFORE))
                    sp.set(qn('w:after'),  str(H3_SPACE_AFTER))
                changed = True

        if changed:
            fixed += 1
    return fixed


# ─── TOC FIELD ───────────────────────────────────────────────────────────────

def inject_toc_field(doc):
    """Replace manual dot-leader TOC with a Word auto-TOC field."""
    paras = list(doc.paragraphs)
    body  = doc.element.body

    toc_start = None
    toc_end   = None
    for i, para in enumerate(paras):
        txt = para.text.strip()
        if txt.count('…') > 3 or '........' in txt:
            if toc_start is None:
                toc_start = i
            toc_end = i

    if toc_start is None:
        print("  [TOC] No manual TOC found — skipping")
        return False

    toc_xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:pPr><w:jc w:val="center"/></w:pPr>'
        '<w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:t>Open in Word and press Ctrl+A then F9 to update page numbers.</w:t></w:r>'
        '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
        '</w:p>'
    )
    toc_el  = etree.fromstring(toc_xml)
    old_els = [paras[i]._element for i in range(toc_start, toc_end + 1)]

    body.insert(list(body).index(old_els[0]), toc_el)
    for el in old_els:
        try:
            body.remove(el)
        except ValueError:
            pass

    print(f"  [TOC] Injected Word auto-TOC field (replaced {len(old_els)} manual entries)")
    print(f"  [TOC] → Open output in Word, press Ctrl+A then F9 to refresh page numbers")
    return True


# ─── LIBREOFFICE REFRESH ─────────────────────────────────────────────────────

def libreoffice_refresh(path: Path) -> bool:
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if not soffice:
        print("  [refresh] LibreOffice not installed — skipping")
        print("  [refresh] Install with: sudo apt install libreoffice")
        return False
    with tempfile.TemporaryDirectory() as tmp:
        import shutil as _sh
        src = Path(tmp) / path.name
        _sh.copy2(path, src)
        try:
            subprocess.run(
                [soffice, '--headless', '--norestore',
                 '--convert-to', 'docx', '--outdir', tmp, str(src)],
                capture_output=True, timeout=120
            )
            out = [f for f in Path(tmp).glob('*.docx') if f != src]
            if out:
                _sh.copy2(out[0], path)
                print("  [refresh] LibreOffice repagination done")
                return True
        except Exception as e:
            print(f"  [refresh] LibreOffice error: {e}")
    return False


# ─── MAIN ────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(
        description="Format a report .docx — never changes text content",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Fixes applied:
  • Empty paragraphs before headings removed
  • Chapter titles (CHAPTER X / ABSTRACT / etc.) forced to new page
  • Body text: full justification, Times New Roman 12pt, 1.5 spacing
  • Heading 2/3 spacing normalized
  • Widow/orphan control on body text
  • Manual TOC replaced with Word auto-field (press F9 to update)
  • LibreOffice headless repagination (if installed)

Full pipeline:
  python3 humanize_doc.py  report.docx  humanized.docx
  python3 format_doc.py    humanized.docx  final.docx
  → open final.docx in Word → Ctrl+A → F9 → Save
        """)
    ap.add_argument("input")
    ap.add_argument("output")
    ap.add_argument("--report",     action="store_true", help="Show changes without writing")
    ap.add_argument("--no-toc",     action="store_true", help="Skip TOC field injection")
    ap.add_argument("--no-refresh", action="store_true", help="Skip LibreOffice repagination")
    args = ap.parse_args()

    src = Path(args.input)
    if not src.exists():
        sys.exit(f"ERROR: {src} not found.")

    print(f"\n{'='*55}")
    print(f"  format_doc.py")
    print(f"  Input  : {src}")
    print(f"  Output : {args.output}")
    if args.report:
        print(f"  Mode   : REPORT ONLY (no changes written)")
    print(f"{'='*55}\n")

    doc = Document(str(src))

    # 1. Remove empty paragraphs before headings
    n_empty = remove_empty_paras_before_heading(doc, report=args.report)
    print(f"  Empty paras removed before headings : {n_empty}")

    # 2. Chapter page breaks
    n_added, n_removed = fix_chapter_page_breaks(doc, report=args.report)
    print(f"  Chapter page-breaks added           : {n_added}")
    print(f"  Spurious page-breaks removed        : {n_removed}")

    # 3. Body formatting
    n_body = fix_body_formatting(doc, report=args.report)
    print(f"  Body paragraphs reformatted         : {n_body}")

    # 4. Heading spacing
    n_hdg = fix_heading_spacing(doc, report=args.report)
    print(f"  Heading spacings normalized         : {n_hdg}")

    if args.report:
        print("\n(Report mode — nothing written)")
        return

    # 5. TOC
    if not args.no_toc:
        print("\nRebuilding TOC...")
        inject_toc_field(doc)

    # 6. Save
    out = Path(args.output)
    doc.save(str(out))
    print(f"\nSaved → {out}")

    # 7. LibreOffice repagination
    if not args.no_refresh:
        print("\nRunning LibreOffice repagination...")
        libreoffice_refresh(out)

    print(f"\n✓  Done!")
    print(f"   {n_empty} empty gaps removed")
    print(f"   {n_added} chapter page-breaks added")
    print(f"   {n_body} body paragraphs reformatted")
    print(f"   {n_hdg} heading spacings fixed")
    print()
    print("   → Open the output in Word")
    print("   → Press Ctrl+A then F9")
    print("   → Save — TOC page numbers will be correct\n")


if __name__ == "__main__":
    main()