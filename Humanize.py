#!/usr/bin/env python3
"""
humanize_doc.py — Humanize a .docx via humanizeai.pro

Usage:
    python3 humanize_doc.py input.docx output.docx
    python3 humanize_doc.py input.docx output.docx --concurrency 5
    python3 humanize_doc.py input.docx output.docx --dry-run
    python3 humanize_doc.py input.docx output.docx --visible --concurrency 1

Requires:
    pip install patchright python-docx
    python3 -m patchright install chromium
"""

import argparse
import asyncio
import copy
import re
import sys
import time
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from patchright.async_api import async_playwright, TimeoutError as PWTimeout
    BROWSER_LIB = "patchright"
except ImportError:
    from playwright.async_api import async_playwright, TimeoutError as PWTimeout
    BROWSER_LIB = "playwright (install patchright to bypass Cloudflare)"


# ─── CONFIG ──────────────────────────────────────────────────────────────────

SITE_URL          = "https://www.humanizeai.pro/"
PAGE_LOAD_TIMEOUT = 60_000
MAX_WAIT_SECS     = 180
POLL_INTERVAL     = 2.0
CHUNK_WORDS       = 175
CONCURRENCY       = 5

SKIP_STYLE_KEYWORDS = [
    'heading', 'title', 'caption', 'toc ', 'table of', 'code',
    'figure', 'list bullet', 'list number', 'footnote', 'header', 'footer',
    'bibliography', 'index',
]
MIN_WORDS = 15

BOILERPLATE_MARKERS = [
    "Start using our AI Humanizer",
    "We bypass ALL detectors",
    "Introducing Humanize AI",
    "Humanize AI stands out",
    "Content Creators & Writers",
    "Marketing Professionals",
    "We use cookies to ensure",
    "humanizeai.pro. All rights reserved",
    "Feature Request & Issue Report",
    "reCAPTCHA and the Google",
    "A2: The tool uses advanced",
    "Anyone with a Story to Tell",
]

LINE_NOISE_RE = [
    re.compile(r"^\d+\s+words?\s*$", re.I),
    re.compile(r"^Free\s+Standard\s+Academic", re.I),
    re.compile(r"^Ultra\s+run$", re.I),
    re.compile(r"^Expand\s+More$", re.I),
    re.compile(r"^Humanize\s+AI\s*$", re.I),
    re.compile(r"^Check\s+for\s+AI", re.I),
    re.compile(r"^Expected\s+time", re.I),
    re.compile(r"added\s+change\s*$", re.I),
]

USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:125.0) Gecko/20100101 Firefox/125.0",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 14_4) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.4.1 Safari/605.1.15",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
]


# ─── DOCX HELPERS ────────────────────────────────────────────────────────────

def should_skip(para):
    txt = para.text.strip()
    if not txt:
        return True
    if para._element.findall('.//' + qn('w:drawing')):
        return True
    if (para._element.findall('.//' + qn('m:oMath')) or
            para._element.findall('.//' + qn('m:oMathPara'))):
        return True
    style = para.style.name.lower()
    if any(kw in style for kw in SKIP_STYLE_KEYWORDS):
        return True
    words = txt.split()
    if len(words) < MIN_WORDS:
        return True
    if not any(c in txt for c in ['.', ',']):
        return True
    if txt.count('…') > 2 or txt.count('.....') > 2:
        return True
    if re.match(r'^\[\d+\]', txt):
        return True
    return False


def extract_paragraphs(doc):
    return [(idx, para, para.text.strip())
            for idx, para in enumerate(doc.paragraphs)
            if not should_skip(para)]


def chunk_paragraphs(para_list, max_words):
    chunks, current, current_words = [], [], 0
    for item in para_list:
        wc = len(item[2].split())
        if wc > max_words:
            if current:
                chunks.append(current)
                current, current_words = [], 0
            chunks.append([item])
        elif current_words + wc > max_words:
            chunks.append(current)
            current, current_words = [item], wc
        else:
            current.append(item)
            current_words += wc
    if current:
        chunks.append(current)
    return chunks


def replace_paragraph_text(para, new_text):
    runs = para.runs
    if not runs:
        new_r = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = new_text
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_r.append(new_t)
        para._element.append(new_r)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


# ─── OUTPUT CLEANING ─────────────────────────────────────────────────────────

def clean_output(raw: str, original: str) -> str:
    cut = len(raw)
    for marker in BOILERPLATE_MARKERS:
        idx = raw.lower().find(marker.lower())
        if 0 <= idx < cut:
            cut = idx
    raw = raw[:cut]

    lines = raw.split('\n')
    cleaned = []
    for line in lines:
        ls = line.strip()
        if not ls:
            continue
        if any(rx.search(ls) for rx in LINE_NOISE_RE):
            continue
        cleaned.append(ls)

    result = '\n\n'.join(cleaned).strip()
    if not result or len(result.split()) < len(original.split()) * 0.3:
        return ""
    return result


# ─── JAVASCRIPT ──────────────────────────────────────────────────────────────

# Step 1: Before clicking Humanize, snapshot ALL text in the right half.
# Returns a string we can compare against later.
JS_SNAPSHOT_RIGHT = """
() => {
    const vw = window.innerWidth;
    const parts = [];
    for (const el of document.querySelectorAll('div, p, section')) {
        const r = el.getBoundingClientRect();
        if (r.left < vw * 0.4) continue;
        if (r.width < 80 || r.height < 20) continue;
        const txt = (el.innerText || '').trim();
        if (txt.length > 20) parts.push(txt.slice(0, 100));
    }
    return parts.join('|||');
}
"""

# Step 2: After clicking, poll until right-side content has CHANGED from snapshot.
# Returns the new text or '' if nothing changed yet.
JS_GET_CHANGED = """
(args) => {
    const snapshot   = args.snapshot;      // text state before clicking
    const inputStart = args.inputStart;    // first 40 chars of input (to exclude left panel)
    const vw = window.innerWidth;
    const inputLow = inputStart.toLowerCase();

    // Find the right-panel element whose text has CHANGED since snapshot
    // and is NOT the input text
    const candidates = [];
    for (const el of document.querySelectorAll('div, p, section')) {
        const r = el.getBoundingClientRect();
        if (r.left < vw * 0.4) continue;
        if (r.width < 150 || r.height < 60) continue;

        const txt = (el.innerText || '').trim();
        if (txt.length < 60) continue;

        // Must NOT be the input
        if (txt.toLowerCase().startsWith(inputLow)) continue;

        // Must NOT have been in the snapshot (i.e. it's NEW content)
        const snip = txt.slice(0, 100);
        if (snapshot.includes(snip)) continue;

        // Prefer elements with <a> links (the colored output words)
        const links = el.querySelectorAll('a').length;
        const area  = r.width * r.height;
        candidates.push({ txt, links, area });
    }

    if (candidates.length === 0) return '';

    // Sort: most links first, smallest area (most specific)
    candidates.sort((a, b) => {
        if (b.links !== a.links) return b.links - a.links;
        return a.area - b.area;
    });

    return candidates[0].txt;
}
"""


# ─── SINGLE CHUNK PROCESSOR ──────────────────────────────────────────────────

async def process_chunk(ctx, chunk_text: str, chunk_id: int, probe: bool) -> str:
    page = await ctx.new_page()
    try:
        # 1. Load page
        print(f"  [chunk {chunk_id:03d}] Loading...")
        await page.goto(SITE_URL, wait_until="domcontentloaded", timeout=PAGE_LOAD_TIMEOUT)
        await asyncio.sleep(3)

        # Scroll editor into viewport (required for lazy render)
        await page.evaluate("window.scrollTo(0, 350)")
        await asyncio.sleep(0.5)

        # Dismiss cookie banner
        for label in ["Ok", "OK", "Accept", "Accept all"]:
            try:
                btn = page.get_by_role("button", name=label)
                if await btn.count() > 0:
                    await btn.first.click(timeout=2000)
                    break
            except Exception:
                pass

        # 2. Find textarea
        textarea = None
        for sel in ["textarea.chakra-textarea", "textarea"]:
            try:
                await page.wait_for_selector(sel, timeout=15_000)
                el = await page.query_selector(sel)
                if el and await el.is_visible():
                    textarea = sel
                    break
            except Exception:
                continue

        if not textarea:
            await page.screenshot(path=f"/tmp/chunk_{chunk_id}_no_input.png")
            print(f"  [chunk {chunk_id:03d}] ⚠  No textarea")
            return chunk_text

        # 3. SNAPSHOT right side BEFORE clicking (captures static marketing divs)
        try:
            snapshot = await page.evaluate(JS_SNAPSHOT_RIGHT)
        except Exception:
            snapshot = ""

        # 4. Fill and click Humanize
        await page.fill(textarea, "")
        await asyncio.sleep(0.2)
        await page.fill(textarea, chunk_text)
        await asyncio.sleep(0.8)

        clicked = False
        for label in ["Humanize", "humanize"]:
            try:
                btn = page.get_by_role("button", name=label)
                if await btn.count() > 0:
                    await btn.first.scroll_into_view_if_needed()
                    await btn.first.click(timeout=10_000)
                    clicked = True
                    break
            except Exception:
                pass

        if not clicked:
            await page.screenshot(path=f"/tmp/chunk_{chunk_id}_no_btn.png")
            print(f"  [chunk {chunk_id:03d}] ⚠  Humanize button not found")
            return chunk_text

        print(f"  [chunk {chunk_id:03d}] {len(chunk_text.split())}w → waiting for output...")

        # 5. Poll until right side has NEW content (different from snapshot)
        input_start = chunk_text.strip()[:40].lower()
        deadline    = time.time() + MAX_WAIT_SECS
        result      = ""
        poll_count  = 0

        while time.time() < deadline:
            poll_count += 1
            await asyncio.sleep(POLL_INTERVAL)

            try:
                raw = await page.evaluate(JS_GET_CHANGED, {
                    "snapshot":   snapshot,
                    "inputStart": input_start,
                })
            except Exception as e:
                if probe:
                    print(f"    poll#{poll_count} error: {e}")
                raw = ""

            if probe:
                print(f"    [chunk {chunk_id:03d}] poll#{poll_count}: {len(raw)}chars — {raw[:80]!r}")

            if raw and len(raw.strip()) > 40:
                cleaned = clean_output(raw, chunk_text)
                if cleaned:
                    result = cleaned
                    break

        # 6. Done
        if not result:
            await page.screenshot(path=f"/tmp/chunk_{chunk_id}_fail.png")
            print(f"  [chunk {chunk_id:03d}] ⚠  No output after {poll_count} polls — screenshot: /tmp/chunk_{chunk_id}_fail.png")
            return chunk_text

        print(f"  [chunk {chunk_id:03d}] ✓  {len(result.split())}w (after {poll_count} polls)")
        return result

    except PWTimeout as e:
        print(f"  [chunk {chunk_id:03d}] TIMEOUT: {str(e).splitlines()[0]}")
        return chunk_text
    except Exception as e:
        print(f"  [chunk {chunk_id:03d}] ERROR: {e}")
        return chunk_text
    finally:
        try:
            await page.close()
        except Exception:
            pass


# ─── WORKER + ORCHESTRATOR ───────────────────────────────────────────────────

async def worker(browser, queue, results, probe, worker_id):
    ctx = await browser.new_context(
        user_agent=USER_AGENTS[worker_id % len(USER_AGENTS)],
        viewport={"width": 1280 + worker_id * 13, "height": 900},
        locale="en-US",
        timezone_id="America/New_York",
    )
    await asyncio.sleep(worker_id * 2.0)
    try:
        while True:
            try:
                chunk_id, chunk_text = queue.get_nowait()
            except asyncio.QueueEmpty:
                break
            results[chunk_id] = await process_chunk(ctx, chunk_text, chunk_id, probe)
            queue.task_done()
    except Exception as e:
        print(f"  [worker {worker_id}] fatal: {e}")
    finally:
        try:
            await ctx.close()
        except Exception:
            pass


async def run_all(chunks_text, concurrency, headless, probe):
    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=headless,
            args=["--no-sandbox", "--disable-dev-shm-usage"],
        )
        try:
            queue   = asyncio.Queue()
            for i, txt in enumerate(chunks_text):
                await queue.put((i, txt))
            results = [None] * len(chunks_text)
            n       = min(concurrency, len(chunks_text))
            print(f"  Spawning {n} isolated browser contexts...\n")
            await asyncio.gather(*[
                worker(browser, queue, results, probe, wid)
                for wid in range(n)
            ])
        finally:
            await browser.close()
    return results


# ─── DOCUMENT RECONSTRUCTION ─────────────────────────────────────────────────

def reconstruct_doc(doc, chunks, results, chunk_texts):
    for chunk, humanized, orig_text in zip(chunks, results, chunk_texts):
        if not humanized or humanized.strip() == orig_text.strip():
            continue
        n  = len(chunk)
        hp = [p.strip() for p in humanized.split("\n\n") if p.strip()]
        if not hp:
            continue
        if len(hp) == n:
            pairs = list(zip(chunk, hp))
        elif len(hp) > n:
            pairs = list(zip(chunk, hp[:n-1]))
            pairs.append((chunk[-1], " ".join(hp[n-1:])))
        else:
            pairs = [(chunk[i], hp[i] if i < len(hp) else hp[-1]) for i in range(n)]
        for (_, para, orig), new_text in pairs:
            if new_text and new_text.strip() != orig.strip():
                replace_paragraph_text(para, new_text)


# ─── MAIN ────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(
        description="Humanize a .docx via humanizeai.pro",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python3 humanize_doc.py report.docx out.docx
  python3 humanize_doc.py report.docx out.docx --concurrency 5
  python3 humanize_doc.py report.docx out.docx --dry-run
  python3 humanize_doc.py report.docx out.docx --visible --concurrency 1   (debug)
  python3 humanize_doc.py report.docx out.docx --probe                     (verbose polls)
        """)
    ap.add_argument("input")
    ap.add_argument("output")
    ap.add_argument("--concurrency", type=int, default=CONCURRENCY)
    ap.add_argument("--chunk-words", type=int, default=CHUNK_WORDS)
    ap.add_argument("--dry-run",  action="store_true")
    ap.add_argument("--visible",  action="store_true")
    ap.add_argument("--probe",    action="store_true", help="Print every poll attempt")
    args = ap.parse_args()

    src = Path(args.input)
    if not src.exists():
        sys.exit(f"ERROR: {src} not found.")

    print(f"\n{'='*55}")
    print(f"  Browser     : {BROWSER_LIB}")
    print(f"  Input       : {src}")
    print(f"  Output      : {args.output}")
    print(f"  Workers     : {args.concurrency}")
    print(f"  Chunk limit : {args.chunk_words} words")
    print(f"  Max wait    : {MAX_WAIT_SECS}s per chunk")
    print(f"{'='*55}\n")

    doc       = Document(str(src))
    para_list = extract_paragraphs(doc)
    total     = sum(1 for p in doc.paragraphs if p.text.strip())
    print(f"Paragraphs : {total} total | {total-len(para_list)} skipped | {len(para_list)} to humanize")

    chunks = chunk_paragraphs(para_list, args.chunk_words)
    words  = sum(len(p[2].split()) for p in para_list)
    eta    = int((len(chunks) / args.concurrency) * 22)
    m, s   = divmod(eta, 60)
    print(f"Chunks     : {len(chunks)}  ({words} words)  ~{m}m {s}s estimated\n")

    chunk_texts = ["\n\n".join(p[2] for p in ch) for ch in chunks]

    if args.dry_run:
        for i, ct in enumerate(chunk_texts):
            print(f"[Chunk {i:03d}] {len(ct.split())}w | {len(chunks[i])} para(s)")
            print(ct[:200] + ("..." if len(ct) > 200 else ""))
            print("-" * 50)
        return

    t0      = time.time()
    results = asyncio.run(run_all(
        chunk_texts, args.concurrency,
        headless=not args.visible,
        probe=args.probe,
    ))
    elapsed = time.time() - t0

    changed = sum(1 for o, r in zip(chunk_texts, results)
                  if r and r.strip() != o.strip())
    print(f"\n  {changed}/{len(chunks)} chunks humanized in {elapsed:.0f}s")
    print("  Rebuilding document...")
    reconstruct_doc(doc, chunks, results, chunk_texts)
    doc.save(args.output)
    m2, s2 = divmod(int(elapsed), 60)
    print(f"\n✓  Saved → {args.output}  ({m2}m {s2}s)\n")


if __name__ == "__main__":
    main()